"""
catalog/utils/docx_generator.py
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Converts BrainzAcademy AI lesson-note markdown to a polished .docx file.

Handles:
  - Headings (#, ##, ###)
  - Bold / italic inline runs
  - Bullet lists  (- / *)
  - Numbered lists (1.)
  - Markdown tables
  - Blockquotes (> lines) — styled as tip/warning callouts
  - Horizontal rules (---)
  - SVG blocks          → placeholder paragraph
  - Mermaid code fences → placeholder paragraph
  - LaTeX math          → plain text (delimiters stripped)
  - Plain paragraphs

Output is an in-memory BytesIO buffer ready for an HttpResponse attachment.
"""

from __future__ import annotations

import io
import re
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


# ── Brand colours ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0B, 0x2D, 0x72)   # #0B2D72
BLUE   = RGBColor(0x09, 0x92, 0xC2)   # #0992C2
MUTED  = RGBColor(0x6B, 0x7A, 0x99)   # mid-grey
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF6, 0xFF)   # table even-row tint
GOLD   = RGBColor(0xF6, 0xE7, 0xBC)   # callout background hint

# ── Regex patterns ────────────────────────────────────────────────────────────
RE_SVG_BLOCK    = re.compile(r'<svg[\s\S]*?</svg>', re.IGNORECASE)
RE_MERMAID      = re.compile(r'```mermaid[\s\S]*?```', re.IGNORECASE)
RE_CODE_FENCE   = re.compile(r'```[\s\S]*?```')
RE_HEADING      = re.compile(r'^(#{1,3})\s+(.*)')
RE_HR           = re.compile(r'^-{3,}$|^\*{3,}$|^_{3,}$')
RE_BULLET       = re.compile(r'^[\*\-]\s+(.*)')
RE_NUMBERED     = re.compile(r'^\d+\.\s+(.*)')
RE_BLOCKQUOTE   = re.compile(r'^>\s*(.*)')
RE_TABLE_ROW    = re.compile(r'^\|')
RE_TABLE_SEP    = re.compile(r'^\|[\s\-\|:]+\|$')
RE_BOLD         = re.compile(r'\*\*(.*?)\*\*')
RE_ITALIC       = re.compile(r'\*(.*?)\*')
RE_LATEX_DISP   = re.compile(r'\\\[[\s\S]*?\\\]')
RE_LATEX_INLINE = re.compile(r'\\\([\s\S]*?\\\)')
RE_BOLD_ITALIC  = re.compile(r'\*\*\*(.*?)\*\*\*')


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _strip_latex(text: str) -> str:
    """Remove LaTeX delimiters, keeping the inner expression as plain text."""
    text = RE_LATEX_DISP.sub(lambda m: m.group(0)[2:-2].strip(), text)
    text = RE_LATEX_INLINE.sub(lambda m: m.group(0)[2:-2].strip(), text)
    return text


def _set_cell_bg(cell, hex_colour: str) -> None:
    """Apply a solid background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)


def _add_paragraph_border_bottom(para, colour: str = '0992C2', size: int = 6) -> None:
    """Add a bottom border to a paragraph (used for HR / section dividers)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    str(size))
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), colour)
    pBdr.append(btm)
    pPr.append(pBdr)


def _add_runs(para, text: str, base_bold: bool = False, base_italic: bool = False) -> None:
    """
    Parse inline markdown bold/italic and add styled runs to *para*.
    LaTeX delimiters are stripped so equations appear as plain text.
    """
    text = _strip_latex(text)

    # Split on ***bold-italic***, **bold**, *italic* tokens
    token_re = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)')
    parts    = token_re.split(text)

    for part in parts:
        if not part:
            continue
        run = para.add_run()
        if part.startswith('***') and part.endswith('***'):
            run.text   = part[3:-3]
            run.bold   = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True or base_bold
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run.text   = part[1:-1]
            run.italic = True or base_italic
        else:
            run.text   = part
            run.bold   = base_bold
            run.italic = base_italic


# ─────────────────────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────────────────────

def _build_document(topic_name: str, subject_name: str) -> Document:
    """Create a new Document with BrainzAcademy styles and a header."""
    doc = Document()

    # ── Page margins (A4, 2 cm all sides) ────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default paragraph font ────────────────────────────────────────────────
    style            = doc.styles['Normal']
    font             = style.font
    font.name        = 'Calibri'
    font.size        = Pt(11)
    font.color.rgb   = RGBColor(0x1A, 0x1A, 0x2E)

    # ── Heading styles ────────────────────────────────────────────────────────
    for level, (pt_size, colour) in enumerate(
        [(16, NAVY), (13, BLUE), (11, NAVY)], start=1
    ):
        h = doc.styles[f'Heading {level}']
        h.font.name      = 'Calibri'
        h.font.size      = Pt(pt_size)
        h.font.bold      = True
        h.font.color.rgb = colour
        h.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        h.paragraph_format.space_after  = Pt(4)

    # ── Document header ───────────────────────────────────────────────────────
    header   = doc.sections[0].header
    hdr_para = header.paragraphs[0]
    hdr_para.clear()
    _add_runs(hdr_para, f'BrainzAcademy  ·  {subject_name}  ·  {topic_name}')
    hdr_para.runs[0].font.size      = Pt(8)
    hdr_para.runs[0].font.color.rgb = MUTED
    hdr_para.alignment              = WD_ALIGN_PARAGRAPH.RIGHT
    _add_paragraph_border_bottom(hdr_para, colour='D1D9E6', size=4)

    # ── Title block ───────────────────────────────────────────────────────────
    title      = doc.add_paragraph()
    title_run  = title.add_run(topic_name)
    title_run.font.name      = 'Calibri'
    title_run.font.size      = Pt(20)
    title_run.font.bold      = True
    title_run.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(2)

    sub      = doc.add_paragraph()
    sub_run  = sub.add_run(f'{subject_name}  ·  Revision Note  ·  BrainzAcademy')
    sub_run.font.size      = Pt(9)
    sub_run.font.color.rgb = MUTED
    sub_run.italic         = True
    sub.paragraph_format.space_after = Pt(10)

    # Title underline rule
    rule = doc.add_paragraph()
    _add_paragraph_border_bottom(rule, colour='0B2D72', size=8)
    rule.paragraph_format.space_after = Pt(6)

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# Block-level renderers
# ─────────────────────────────────────────────────────────────────────────────

def _render_heading(doc: Document, text: str, level: int) -> None:
    text = _strip_latex(RE_BOLD.sub(r'\1', text))   # strip markdown bold inside headings
    doc.add_heading(text, level=level)


def _render_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(2)
    _add_runs(para, text)


def _render_numbered(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.space_after = Pt(2)
    _add_runs(para, text)


def _render_blockquote(doc: Document, text: str) -> None:
    """
    Render a blockquote line (> ...) as an indented callout paragraph.
    ⚠️ Mistake: lines get orange-ish treatment; 💡 Tip: lines get blue.
    """
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.8)
    para.paragraph_format.space_after  = Pt(3)

    # Detect emoji prefix for colour coding
    is_tip     = text.startswith('💡')
    is_mistake = text.startswith('⚠️')

    run            = para.add_run()
    run.font.size  = Pt(10)
    run.font.italic = True

    if is_tip:
        run.font.color.rgb = BLUE
    elif is_mistake:
        run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)   # amber
    else:
        run.font.color.rgb = MUTED

    _add_runs(para, text)

    # Left border on the paragraph
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '0992C2' if is_tip else 'E07B29' if is_mistake else 'AAAAAA')
    pBdr.append(left)
    pPr.append(pBdr)


def _render_hr(doc: Document) -> None:
    para = doc.add_paragraph()
    _add_paragraph_border_bottom(para, colour='0B2D72', size=6)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)


def _render_svg_placeholder(doc: Document) -> None:
    """Diagram placeholder — directs reader to online/PDF version."""
    para            = doc.add_paragraph()
    run             = para.add_run(
        '📊  [Diagram available in the PDF download and online version at brainzacademy.com]'
    )
    run.font.size      = Pt(9)
    run.font.italic    = True
    run.font.color.rgb = MUTED
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.5)


def _render_mermaid_placeholder(doc: Document) -> None:
    """Flowchart placeholder — directs reader to online/PDF version."""
    para = doc.add_paragraph()
    run  = para.add_run(
        '🔀  [Flowchart/diagram available in the PDF download and online version at brainzacademy.com]'
    )
    run.font.size      = Pt(9)
    run.font.italic    = True
    run.font.color.rgb = MUTED
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.5)


def _render_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(5)
    _add_runs(para, text)


def _render_table(doc: Document, rows: List[List[str]]) -> None:
    """
    Render a markdown table.
    rows[0] = header row; rows[1] = separator (skipped); rows[2:] = data rows.
    """
    if len(rows) < 3:
        return

    header_cells = [c.strip() for c in rows[0].strip('|').split('|')]
    data_rows    = [
        [c.strip() for c in r.strip('|').split('|')]
        for r in rows[2:]   # skip separator row
        if r.strip() and not RE_TABLE_SEP.match(r.strip())
    ]

    col_count = len(header_cells)
    if col_count == 0:
        return

    # Total content width in DXA: A4 with 2.5 cm margins ≈ 8720 DXA
    TABLE_WIDTH = 8720
    col_width   = TABLE_WIDTH // col_count

    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = 'Table Grid'

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, '0B2D72')
        cell.width = Pt(col_width)
        para  = cell.paragraphs[0]
        run   = para.add_run(_strip_latex(cell_text))
        run.bold            = True
        run.font.color.rgb  = WHITE
        run.font.size       = Pt(10)
        para.alignment      = WD_ALIGN_PARAGRAPH.LEFT

    # ── Data rows ─────────────────────────────────────────────────────────────
    for r_idx, row_data in enumerate(data_rows):
        row = table.rows[r_idx + 1]
        bg  = 'F0F6FF' if r_idx % 2 == 1 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            cell.width = Pt(col_width)
            para  = cell.paragraphs[0]
            _add_runs(para, _strip_latex(cell_text))
            para.runs[0].font.size = Pt(10) if para.runs else Pt(10)

    # Space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─────────────────────────────────────────────────────────────────────────────
# Main parser
# ─────────────────────────────────────────────────────────────────────────────

def _preprocess(raw: str) -> str:
    """
    Strip SVG and mermaid blocks, replacing them with sentinel placeholders
    so the line-by-line parser can emit the correct Word elements.
    """
    raw = RE_SVG_BLOCK.sub('\n%%SVG_PLACEHOLDER%%\n', raw)
    raw = RE_MERMAID.sub('\n%%MERMAID_PLACEHOLDER%%\n', raw)
    raw = RE_CODE_FENCE.sub('\n%%CODE_PLACEHOLDER%%\n', raw)
    return raw


def generate_lesson_note_docx(
    ai_content: str,
    topic_name: str,
    subject_name: str,
) -> io.BytesIO:
    """
    Convert *ai_content* (raw markdown from the AI lesson-note generator)
    into a formatted .docx and return it as an in-memory BytesIO buffer.

    Usage::
        buf = generate_lesson_note_docx(note.ai_content, topic.name, subject.name)
        response = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    """
    doc = _build_document(topic_name, subject_name)

    raw   = _preprocess(ai_content)
    lines = raw.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Placeholders ──────────────────────────────────────────────────────
        if line.strip() == '%%SVG_PLACEHOLDER%%':
            _render_svg_placeholder(doc)
            i += 1
            continue

        if line.strip() == '%%MERMAID_PLACEHOLDER%%':
            _render_mermaid_placeholder(doc)
            i += 1
            continue

        if line.strip() == '%%CODE_PLACEHOLDER%%':
            # Generic code block — skip silently (not relevant to text notes)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if RE_HR.match(line.strip()):
            _render_hr(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        m = RE_HEADING.match(line)
        if m:
            level = min(len(m.group(1)), 3)
            _render_heading(doc, m.group(2), level)
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────────────────────
        if RE_TABLE_ROW.match(line):
            table_lines = []
            while i < len(lines) and (RE_TABLE_ROW.match(lines[i]) or RE_TABLE_SEP.match(lines[i].strip())):
                table_lines.append(lines[i])
                i += 1
            _render_table(doc, table_lines)
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        m = RE_BULLET.match(line)
        if m:
            _render_bullet(doc, m.group(1))
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        m = RE_NUMBERED.match(line)
        if m:
            _render_numbered(doc, m.group(1))
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        m = RE_BLOCKQUOTE.match(line)
        if m:
            _render_blockquote(doc, m.group(1))
            i += 1
            continue

        # ── Plain paragraph ───────────────────────────────────────────────────
        _render_paragraph(doc, line)
        i += 1

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para            = doc.add_paragraph()
    footer_run             = footer_para.add_run(
        'Generated by BrainzAcademy AI · brainzacademy.com · '
        'Review before sharing with students.'
    )
    footer_run.font.size      = Pt(8)
    footer_run.font.italic    = True
    footer_run.font.color.rgb = MUTED
    footer_para.alignment     = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_border_bottom(footer_para, colour='0B2D72', size=4)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer