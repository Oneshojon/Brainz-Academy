import re

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import generics
from .models import Question, Theme
from .serializers import (
    SubjectSerializer, TopicSerializer, ThemeSerializer,
    ExamBoardSerializer, QuestionSerializer, QuestionListSerializer,
)
from .permissions import IsTeacher
from rest_framework.permissions import AllowAny, IsAuthenticated
from catalog.subscription_access import check_test_builder_access
from catalog.models import FreeUsageTracker
from django.db import transaction
from catalog.models import SavedTest, SavedTestQuestion, FreeUsageTracker, UserSubscription

# File
from django.http import HttpResponse
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
from datetime import date

from django.http import JsonResponse
from .models import ExamSeries

SITTING_LABELS = {
    'MAY_JUNE': 'May/June',
    'NOV_DEC':  'Nov/Dec',
    'MOCK':     'Mock',
    'OTHER':    'Other',
}




_FIRST_P_RE   = re.compile(r'^(<p[^>]*>)', re.IGNORECASE)
_BLOCK_TAG_RE = re.compile(r'^<(table|figure|img|div|ul|ol)', re.IGNORECASE)

from catalog.cache_utils import (
    get_all_subjects, get_all_boards, get_available_sittings, get_subjects_with_question_counts, get_themes_for_subject,
    get_topics_for_subject, get_topics_for_theme, get_available_years,
    get_feature_flags, invalidate_subject_caches, invalidate_feature_flags,
    get_leaderboard, get_boards_with_question_counts, get_platform_settings
)


class AvailableSittingsView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        subject_id = request.query_params.get('subject')
        board_id   = request.query_params.get('board')
        year       = request.query_params.get('year')

        sittings = get_available_sittings(subject_id, board_id, year)
        data = [
            {'value': s, 'label': SITTING_LABELS.get(s, s)}
            for s in sittings if s
        ]
        return Response({'sittings': data})


class FeatureFlagsView(APIView):
    """
    GET /api/catalog/feature-flags/
    Returns all feature flags as a flat dict {key: is_enabled}.
    Used by the React frontend to check which modes are enabled.
    """
    permission_classes = [AllowAny]

    def get(self, request):
        from catalog.cache_utils import get_feature_flags
        return Response(get_feature_flags())

class SubjectListView(generics.ListAPIView):
    permission_classes = [AllowAny]
    serializer_class   = SubjectSerializer
 
    def get_queryset(self):
        return get_all_subjects()
 
 
class SubjectListView(generics.ListAPIView):
    permission_classes = [AllowAny]
    serializer_class   = SubjectSerializer

    def get_queryset(self):
        return get_subjects_with_question_counts()   
    
class TopicListView(generics.ListAPIView):
    """Returns topics filtered by subject."""
    permission_classes = [AllowAny]
    serializer_class   = TopicSerializer
 
    def get_queryset(self):
        subject_id = self.request.query_params.get('subject')
        if subject_id:
            return get_topics_for_subject(subject_id)
        return []

class TopicsByThemeView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        theme_id = request.query_params.get('theme')
        exam_board_id = request.query_params.get('exam_board')

        if not theme_id:
            return Response([])

        from catalog.cache_utils import get_topics_for_theme_with_counts
        return Response(
            get_topics_for_theme_with_counts(theme_id, exam_board_id)
        )

class ThemeListView(generics.ListAPIView):
    """Returns themes for a subject."""
    permission_classes = [AllowAny]
    serializer_class   = ThemeSerializer
 
    def get_queryset(self):
        subject_id = self.request.query_params.get('subject')
        if subject_id:
            return get_themes_for_subject(subject_id)
        return []    

class AvailableYearsView(APIView):
    permission_classes = [AllowAny]
 
    def get(self, request):
        subject_id    = request.query_params.get('subject')
        exam_board_id = request.query_params.get('exam_board')
        years = get_available_years(subject_id, exam_board_id)
        return Response({'years': years})

class ExamBoardListView(generics.ListAPIView):
    permission_classes = [AllowAny]
    serializer_class   = ExamBoardSerializer

    def get_queryset(self):
        return get_boards_with_question_counts()

class TestBuilderAccessView(APIView):
    """
    GET /api/catalog/test-builder-access/
    Returns the current user\'s test builder permissions.
    The React frontend calls this on load to show/hide options.
    """
    permission_classes = [IsAuthenticated]
 
    def get(self, request):
        access = check_test_builder_access(request.user)
        return Response(access)
 
 
class GenerateQuestionsView(APIView):
    """Updated to enforce free tier limits before generating."""
    permission_classes = [IsAuthenticated]
 
    def post(self, request):
        from catalog.subscription_access import check_test_builder_access
        from catalog.models import FreeUsageTracker
 
        access = check_test_builder_access(request.user)
 
        if not access['allowed']:
            return Response({'error': access['reason']}, status=403)
 
        subject_id    = request.data.get('subject')
        topic_ids     = request.data.get('topics', [])
        exam_board_id = request.data.get('exam_board')
        years         = request.data.get('years', [])
        question_type = request.data.get('question_type', '')
        requested_num = int(request.data.get('num_questions', 15))
 
        # Cap for free users
        num_questions = min(requested_num, access['max_questions'])
 
        qs = Question.objects.all()
        if subject_id:
            qs = qs.filter(subject_id=subject_id)
        if topic_ids:
            qs = qs.filter(topics__id__in=topic_ids).distinct()
        if exam_board_id:
            qs = qs.filter(exam_series__exam_board_id=exam_board_id)
        if years:
            qs = qs.filter(exam_series__year__in=years)
        if question_type:
            qs = qs.filter(question_type=question_type)

        difficulty = request.data.get('difficulty', '')
        if difficulty:
            qs = qs.filter(difficulty=difficulty)

        sitting = request.data.get('sitting', '')
        if sitting:
            qs = qs.filter(exam_series__sitting=sitting)

        qs = qs.select_related(
            'subject',
            'exam_series',
            'exam_series__exam_board',
        ).prefetch_related(
            'choices',
            'topics',
            'theory_answer',
        )

        # Always randomise for free trial; paid can choose ordering later
        if access['is_free']:
            qs = qs.order_by('?')

        questions = list(qs[:num_questions])
 
       

 
        if not questions:
            return Response({'error': 'No questions found for your selection.'}, status=404)
 
        # Increment trial counter for free teachers
        if access['is_free']:
            tracker, _ = FreeUsageTracker.objects.get_or_create(user=request.user)
            tracker.increment_test_builder_trial()
 
        serializer = QuestionSerializer(questions, many=True)
        return Response({
            'questions':        serializer.data,
            'total':            len(questions),
            'is_free':          access['is_free'],
            'pdf_only':         access['pdf_only'],
            'trials_remaining': access['trials_remaining'] - (1 if access['is_free'] else 0),
        })
    

class QuestionDetailView(APIView):
    """
    GET /api/catalog/questions/<id>/
    Returns full question detail including choices and theory answer.
    Used by the test builder preview panel when teacher clicks a question.
    No N+1 — uses select_related + prefetch_related.
    """
    permission_classes = [IsAuthenticated]
 
    def get(self, request, pk):
        try:
            question = (
                Question.objects
                .select_related(
                    'subject',
                    'exam_series',
                    'exam_series__exam_board',
                )
                .prefetch_related(
                    'choices',
                    'topics',
                    'theory_answer',
                )
                .get(pk=pk)
            )
        except Question.DoesNotExist:
            return Response({'error': 'Question not found'}, status=404)
 
        serializer = QuestionSerializer(question)
        return Response(serializer.data)
    

import re
import os
import subprocess
import tempfile
import logging
from datetime import date
from django.db.models import Prefetch
from catalog.models import Choice, Topic

logger = logging.getLogger(__name__)

# ── Compiled patterns at module level — not recompiled on every call ──────────
_BR_RE          = re.compile(r'<br\s*/?>', re.IGNORECASE)
_TAG_RE         = re.compile(r'<[^>]+>')
_ENTITY_MAP     = {'&amp;': '&', '&lt;': '<', '&gt;': '>', '&nbsp;': ' '}
_ENTITY_RE      = re.compile(r'&(?:amp|lt|gt|nbsp);')


def _strip_html(text):
    """Strip HTML tags and decode basic entities."""
    if not text:
        return ''
    text = _BR_RE.sub('\n', text)
    text = _TAG_RE.sub('', text)
    text = _ENTITY_RE.sub(lambda m: _ENTITY_MAP[m.group()], text)
    return text.strip()


def _get_image_bytes(question):
    """Returns (bytes, ext) or (None, None)."""
    try:
        if question.image and question.image.name:
            question.image.open('rb')
            data = question.image.read()
            question.image.close()
            ext = os.path.splitext(question.image.name)[1].lower().lstrip('.') or 'png'
            return data, ext
    except Exception:
        logger.exception(f'Failed to read image for question {question.pk}')
    return None, None


def _prefetch_questions(questions):
    """
    Apply select_related and prefetch_related to a questions queryset.
    Call this in the view before passing questions to any generator.
    Eliminates all N+1s for subject, exam_series, exam_board, choices, topics.
    """
    return (
        questions
        .select_related('subject', 'exam_series__exam_board')
        .prefetch_related(
            Prefetch('choices', queryset=Choice.objects.order_by('label')),
            'topics',
        )
    )


def _header_info(questions, title):
    """
    Build header metadata from questions.
    Requires questions to have select_related('subject', 'exam_series__exam_board') applied.
    All FK traversals hit the prefetch cache — zero extra queries.
    """
    subjects, boards, years = set(), set(), set()
    for q in questions:
        if q.subject_id:                        # use _id to avoid any FK hit
            subjects.add(q.subject.name)
        if q.exam_series_id:
            if q.exam_series.exam_board_id:
                boards.add(q.exam_series.exam_board.name)
            if q.exam_series.year:
                years.add(str(q.exam_series.year))
    return {
        'subject': ', '.join(sorted(subjects)) or title,
        'exam':    ', '.join(sorted(boards))   or '—',
        'year':    ', '.join(sorted(years))    or str(date.today().year),
    }

 
 
class QuestionsByTopicView(APIView):
    """
    GET /api/catalog/questions-by-topic/?topic=<id>&exam_board=<id>
    Returns lightweight question list for the test builder list panel.
    Full detail is fetched separately via QuestionDetailView when previewing.
    """
    permission_classes = [IsAuthenticated]
 
    def get(self, request):
        topic_id      = request.query_params.get('topic')
        exam_board_id = request.query_params.get('exam_board')

        if not topic_id:
            return Response({'error': 'topic is required'}, status=400)

        qs = (
            Question.objects
            .filter(topics__id=topic_id)
            .select_related('subject', 'exam_series', 'exam_series__exam_board')
            .prefetch_related('choices', 'topics', 'theory_answer')
            .order_by('-exam_series__year', 'question_number')
        )

        if exam_board_id:
            qs = qs.filter(exam_series__exam_board_id=exam_board_id)

        serializer = QuestionSerializer(qs, many=True)
        return Response(serializer.data)


# ════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR  (mirrors docx structure exactly)
# ════════════════════════════════════════════════════════════════════════════

# ── Shared HTML builder — used by both DOCX and PDF generators ───────────────

def _build_question_html(questions, title, include_answers=False, marks_map=None, total_marks=0, fmt='pdf'):
    hdr       = _header_info(questions, title)
    copy_line = "Copy: Student" if not include_answers else "Copy: Teacher (With Answers & Topics)"

    # ── Pre-compiled table style regexes ──────────────────────────────────────
    _TABLE_RE = re.compile(r'<table([^>]*)>', re.IGNORECASE)
    _CELL_RE  = re.compile(r'<(td|th)([^>]*)>', re.IGNORECASE)

    def _inject_table_styles(html):
        """Inject inline styles into HTML tables for pandoc/pdf rendering."""
        def replace_table(m):
            attrs = m.group(1)
            if 'border-collapse' in attrs or 'marks-row' in attrs:
                return m.group(0)
            return f'<table{attrs} style="border-collapse:collapse;width:100%;margin:0.5em 0">'

        def replace_cell(m):
            tag   = m.group(1)
            attrs = m.group(2)
            if 'border:' in attrs:
                return m.group(0)
            base = 'border:1px solid #999;padding:6px 10px;vertical-align:top;'
            if tag.lower() == 'th':
                base += 'background:#f0f0f0;font-weight:bold;'
            return f'<{tag}{attrs} style="{base}">'

        html = _TABLE_RE.sub(replace_table, html)   # ← these two lines were missing
        html = _CELL_RE.sub(replace_cell, html)
        return html

    def _lines(images_dir):
        yield (
            '<html><head><meta charset="utf-8"><style>\n'
            'table { border-collapse: collapse; width: 100%; margin: 1em 0; }\n'
            'td, th { border: 1px solid #999; padding: 6px 10px; vertical-align: top; }\n'
            'th { background: #f0f0f0; font-weight: bold; }\n'
            'tr.odd  { background: #ffffff; }\n'
            'tr.even { background: #f9f9f9; }\n'
            '</style></head><body>\n'
        )
        yield f'<p>Subject: {hdr["subject"]}</p>\n'
        yield f'<p>Exam: {hdr["exam"]}</p>\n'
        yield f'<p>Year: {hdr["year"]}</p>\n'
        paper_type = 'Theory' if any(q.question_type == 'THEORY' for q in questions) else 'CBT'
        yield f'<p>Paper Type: {paper_type}</p>\n'
        yield f'<p>{copy_line}</p>\n'
        yield '<p>&nbsp;</p>\n'

        for i, q in enumerate(questions, 1):
            # Resolve marks: custom from Step5 > parsed from file > default 1
            q_marks = (marks_map or {}).get(q.id) or q.marks or 1
            marks_label = f'[{q_marks} mark{"s" if q_marks != 1 else ""}]'

            content = _inject_table_styles(q.content.strip())

            if _FIRST_P_RE.match(content):
                content = _FIRST_P_RE.sub(
                    rf'\1<strong>{i}.</strong>&nbsp;', content, count=1
                )
            elif _BLOCK_TAG_RE.match(content):
                content = f'<p><strong>{i}.</strong></p>\n{content}'
            else:
                content = f'<p><strong>{i}.</strong>&nbsp;{content}</p>'

            # Append marks label right-aligned after question content
            if q.question_type == 'THEORY':
                if fmt == 'pdf':
                    # Use raw LaTeX hfill for reliable right-alignment in xelatex
                    content += '\n<p>\\hfill\\textit{' + marks_label + '}</p>'
                else:
                    content += f'\n<p><em>{marks_label}</em></p>'

            yield content + '\n'

            img_bytes, img_ext = _get_image_bytes(q)
            if img_bytes:
                img_filename = f'q{i}.{img_ext or "png"}'
                with open(os.path.join(images_dir, img_filename), 'wb') as f:
                    f.write(img_bytes)
                yield f'<p><img src="images/{img_filename}" style="max-width:400px"/></p>\n'

            if q.question_type == 'OBJ':
                choices = list(q.choices.all())
                for c in choices:
                    yield f'<p>{c.label}. {c.choice_text}</p>\n'
                if include_answers:
                    correct = next((c for c in choices if c.is_correct), None)
                    if correct:
                        yield f'<p><strong>Answer: {correct.label}</strong></p>\n'

            if include_answers:
                topic_names = [t.name for t in q.topics.all()]
                if topic_names:
                    yield f'<p><strong>Topic: {", ".join(topic_names)}</strong></p>\n'

            yield '<p>&nbsp;</p>\n'

        if total_marks:
            yield '<p>&nbsp;</p>\n'
            if fmt == 'pdf':
                yield f'<p>\\hfill\\textbf{{Total: {total_marks} marks}}</p>\n'
            else:
                yield f'<p style="text-align:right"><strong>Total: {total_marks} marks</strong></p>\n'

        yield '</body></html>\n'    

    return _lines


def _add_table_borders(docx_path):
    """Post-process DOCX to add borders to all tables pandoc generated."""
    doc = Document(docx_path)

    def set_table_borders(table):
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '999999')
            tblBorders.append(border)

        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblBorders)

    for table in doc.tables:
        set_table_borders(table)

    doc.save(docx_path)

def _fix_marks_alignment(docx_path):
    """Right-align marks and total marks paragraphs — post-process after pandoc."""
    marks_re = re.compile(r'^\[(\d+)\s*marks?\]$', re.IGNORECASE)
    total_re = re.compile(r'^Total:\s*\d+\s*marks?$', re.IGNORECASE)
    doc = Document(docx_path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if marks_re.match(text) or total_re.match(text):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(docx_path)

def _generate_docx(questions, title, include_answers=False, marks_map=None, total_marks=0):
    """
    Generate DOCX. Expects _prefetch_questions() applied by caller.
    Returns bytes.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        html_path  = os.path.join(tmpdir, 'input.html')
        docx_path  = os.path.join(tmpdir, 'output.docx')
        images_dir = os.path.join(tmpdir, 'images')
        os.makedirs(images_dir, exist_ok=True)

        html_lines = _build_question_html(questions, title, include_answers, marks_map, total_marks, fmt='docx')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.writelines(html_lines(images_dir))

        result = subprocess.run(
            [
                'pandoc', html_path,
                '-o', docx_path,
                '--from', 'html+tex_math_single_backslash',
                '--to', 'docx',
                '--metadata', f'title={title}',
                '--resource-path', tmpdir,
            ],
            capture_output=True, timeout=60,
            cwd=tmpdir,
        )

        if not os.path.exists(docx_path):
            raise ValueError(
                f'pandoc DOCX failed: rc={result.returncode} '
                f'{result.stderr.decode()[:300]}'
            )

        # ── Post-process: add table borders pandoc stripped ───────────────
        _add_table_borders(docx_path)
        _fix_marks_alignment(docx_path)
        # ──────────────────────────────────────────────────────────────────

        with open(docx_path, 'rb') as f:
            return f.read()

def _generate_pdf(questions, title, include_answers=False, marks_map=None, total_marks=0):
    """
    Generate PDF directly from HTML via xelatex — single pandoc call.
    No DOCX intermediate. Expects _prefetch_questions() applied by caller.
    Returns bytes.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        html_path  = os.path.join(tmpdir, 'input.html')
        pdf_path   = os.path.join(tmpdir, 'output.pdf')
        images_dir = os.path.join(tmpdir, 'images')
        os.makedirs(images_dir, exist_ok=True)

        html_lines = _build_question_html(questions, title, include_answers, marks_map, total_marks, fmt='pdf')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.writelines(html_lines(images_dir))

        result = subprocess.run(
            [
                'pandoc', html_path,
                '-o', pdf_path,
                '--from', 'html+tex_math_single_backslash',
                '--pdf-engine', 'xelatex',
                '--variable', 'geometry:margin=1in',
                '--variable', 'tables=true',
                '--variable', 'colorlinks=true',
                '--resource-path', tmpdir,
            ],
            capture_output=True, timeout=120,
            cwd=tmpdir,
        )

        if not os.path.exists(pdf_path):
            raise ValueError(
                f'PDF failed: rc={result.returncode} '
                f'{result.stderr.decode()[:300]}'
            )

        with open(pdf_path, 'rb') as f:
            return f.read()

# VIEW
# ════════════════════════════════════════════════════════════════════════════

class QuestionDownloadView(APIView):
    """
    POST /api/catalog/questions/download/

    Body JSON:
        question_ids   : [1, 2, 3, ...]
        title          : "Physics WAEC 2020"
        format         : "pdf" | "docx"
        copy_type      : "student" | "teacher"
        builder_mode   : "manual" | "random"          (default: "manual")
        custom_marks   : {"<id>": marks, ...}          (optional, from Step 5)
        saved_test_id  : int | null                    (null = new test, int = update existing)
        total_marks    : int                           (snapshot for SavedTest record)
    """
    permission_classes = [IsTeacher]

    def post(self, request):
        question_ids  = request.data.get('question_ids', [])
        title         = request.data.get('title', 'Question Set')
        fmt           = request.data.get('format', 'pdf').lower()
        copy_type     = request.data.get('copy_type', 'student')
        builder_mode  = request.data.get('builder_mode', 'manual')
        custom_marks  = request.data.get('custom_marks', {})   # {"<id>": int}
        saved_test_id = request.data.get('saved_test_id')      # None = new test
        total_marks   = request.data.get('total_marks', 0)

        # ── Validate inputs ───────────────────────────────────────────────────
        if fmt not in ('docx', 'pdf'):
            return Response(
                {'error': f'Invalid format "{fmt}". Must be "docx" or "pdf".'},
                status=400
            )

        if (
            not isinstance(question_ids, list)
            or not question_ids
            or not all(isinstance(i, int) for i in question_ids)
        ):
            return Response(
                {'error': 'question_ids must be a non-empty list of integers.'},
                status=400
            )

        # ── Access / trial check ──────────────────────────────────────────────
        from catalog.subscription_access import has_subscription
        is_pro = (
            getattr(request.user, 'is_admin', False) or
            has_subscription(request.user, 'TEACHER_PRO')
        )

        tracker = None  # only fetched for free-tier teachers

        # Determine whether this download requires a new trial:
        #   - No saved_test_id → brand new test → always costs a trial
        #   - saved_test_id exists → compare stored question set vs incoming
        #     (as sets, ignoring order) → changed = costs a trial
        existing_test  = None
        questions_changed = True  # default: assume new

        if saved_test_id:
            existing_test = (
                SavedTest.objects
                .filter(pk=saved_test_id, teacher=request.user)
                .first()
            )
            if existing_test:
                stored_ids = set(
                    existing_test.test_questions
                    .values_list('question_id', flat=True)
                )
                incoming_ids = set(question_ids)
                questions_changed = stored_ids != incoming_ids

        needs_trial = not existing_test or questions_changed

        if needs_trial and not is_pro:
            tracker, _ = FreeUsageTracker.objects.get_or_create(user=request.user)
            if not tracker.can_use_test_builder():
                return Response(
                    {
                        'error': (
                            f"You've used all {get_platform_settings().free_test_builder_trials} "
                            f"free trials. Upgrade to Teacher Pro to continue."
                        ),
                        'allowed': False,
                    },
                    status=403
                )

        # ── Free tier: enforce pdf_only ───────────────────────────────────────
        if not is_pro and fmt == 'docx':
            return Response(
                {'error': 'Word (.docx) download requires Teacher Pro. PDF is available on free tier.'},
                status=403
            )

        # ── Query — one trip, no N+1 ──────────────────────────────────────────
        # Preserve the user's chosen order (question_ids order from frontend)
        id_to_pos = {qid: idx for idx, qid in enumerate(question_ids)}
        qs_list = list(
            Question.objects
            .filter(id__in=question_ids)
            .select_related('subject', 'exam_series', 'exam_series__exam_board')
            .prefetch_related('choices', 'topics', 'theory_answer')
        )
        qs_list.sort(key=lambda q: id_to_pos.get(q.id, 9999))
        
        marks_map = {
            q.id: int(custom_marks.get(str(q.id), q.marks or 1))
            for q in qs_list
        }

        if not qs_list:
            return Response({'error': 'No questions found.'}, status=400)

        # ── Generate file ─────────────────────────────────────────────────────
        include_answers = (copy_type == 'teacher')
        safe_title      = title.replace(' ', '_').replace('/', '-')
        copy_label      = 'teacher' if include_answers else 'student'
        filename        = f"{safe_title}_{copy_label}"

        try:
            if fmt == 'docx':
                content  = _generate_docx(qs_list, title,
                                        include_answers=include_answers,
                                        marks_map=marks_map,
                                        total_marks=total_marks)
                ct       = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                response = HttpResponse(content, content_type=ct)
            else:
                content  = _generate_pdf(qs_list, title,
                                        include_answers=include_answers,
                                        marks_map=marks_map,
                                        total_marks=total_marks)
                response = HttpResponse(content, content_type='application/pdf')

            response['Content-Disposition'] = f'attachment; filename="{filename}.{fmt}"'

        except Exception as e:
            logger.exception(f'File generation failed for title="{title}" fmt={fmt}')
            return Response({'error': f'File generation failed: {e}'}, status=500)

        # ── Persist SavedTest record (after successful file generation) ────────
        try:
            with transaction.atomic():
                if existing_test and not questions_changed:
                    # ── Questions unchanged: update metadata only, no trial ────
                    existing_test.title      = title[:255]
                    existing_test.format     = fmt
                    existing_test.copy_type  = copy_type
                    existing_test.total_marks = total_marks
                    existing_test.save(update_fields=[
                        'title', 'format', 'copy_type', 'total_marks', 'updated_at',
                    ])
                    # Update custom_marks in through records (marks may have changed)
                    for stq in existing_test.test_questions.all():
                        new_marks = int(custom_marks.get(str(stq.question_id), stq.custom_marks))
                        if stq.custom_marks != new_marks:
                            stq.custom_marks = new_marks
                            stq.save(update_fields=['custom_marks'])
                    test = existing_test

                else:
                    # ── Questions changed or brand new: create fresh record ────
                    test = SavedTest.objects.create(
                        teacher        = request.user,
                        title          = title[:255],
                        format         = fmt,
                        copy_type      = copy_type,
                        builder_mode   = builder_mode,
                        question_count = len(question_ids),
                        total_marks    = total_marks,
                        # Link to previous version for lineage tracking
                        cloned_from    = existing_test if existing_test else None,
                    )
                    SavedTestQuestion.objects.bulk_create([
                        SavedTestQuestion(
                            saved_test   = test,
                            question_id  = qid,
                            custom_marks = int(custom_marks.get(str(qid), 1)),
                            order        = idx,
                        )
                        for idx, qid in enumerate(question_ids)
                    ])

                    # Consume trial for free-tier teachers
                    if needs_trial and not is_pro:
                        if tracker is None:
                            tracker, _ = FreeUsageTracker.objects.get_or_create(user=request.user)
                        tracker.increment_test_builder_trial()

                # Always return the current test PK so frontend stays in sync
                response['X-Saved-Test-Id'] = str(test.pk)

        except Exception:
            logger.exception(
                f'SavedTest persistence failed for user={request.user.id} '
                f'title="{title}" saved_test_id={saved_test_id}'
            )

        return response