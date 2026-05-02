"""
Unit tests: file parsers.

DOCX parser: _parse_docx(file_bytes) is a private function inside
teacher/views.py — not importable directly. Tests go through the
upload_docx view endpoint instead.

PDF note parser: parse_note_pdf(file_bytes) in catalog/note_pdf_parser.py
— importable and tested directly with mocks.
"""

import pytest
import io
from unittest.mock import patch, MagicMock
from django.urls import reverse


# ---------------------------------------------------------------------------
# DOCX upload — tested through the view (since _parse_docx is private)
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestDocxUploadView:

    def _make_docx_bytes(self, paragraphs):
        """Build a minimal in-memory DOCX from paragraph strings."""
        from docx import Document
        doc = Document()
        for p in paragraphs:
            doc.add_paragraph(p)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def test_upload_docx_requires_login(self, client):
        url = reverse('teacher:upload_docx')
        response = client.get(url)
        assert response.status_code == 302

    def test_upload_docx_page_renders_for_admin(self, client, admin_user):
        client.force_login(admin_user)
        url = reverse('teacher:upload_docx')
        response = client.get(url)
        assert response.status_code == 200

    def test_upload_docx_requires_file(self, client, admin_user):
        """POST with no file should not 500."""
        client.force_login(admin_user)
        url = reverse('teacher:upload_docx')
        response = client.post(url, {})
        assert response.status_code in (200, 302, 400)

    def test_upload_docx_with_valid_file(self, client, admin_user, subject):
        """
        Upload a minimal well-formed DOCX.
        Parser may return 0 questions for a trivial file — we just
        confirm the view handles it without a 500.
        """
        client.force_login(admin_user)
        url = reverse('teacher:upload_docx')
        paragraphs = [
            '1. What is the speed of light?',
            'A. 3×10⁸ m/s',
            'B. 3×10⁶ m/s',
            'C. 3×10⁴ m/s',
            'D. 3×10² m/s',
            'Answer: A',
            'Topic: Waves',
        ]
        docx_buf = self._make_docx_bytes(paragraphs)
        docx_buf.name = 'test.docx'

        response = client.post(url, {
            'subject':    subject.id,
            'exam_board': '',
            'year':       2023,
            'sitting':    'MAY_JUNE',
            'file':       docx_buf,
        })
        assert response.status_code in (200, 302)

    def test_non_admin_cannot_access_upload_docx(self, client, student):
        client.force_login(student)
        url = reverse('teacher:upload_docx')
        response = client.get(url)
        assert response.status_code in (302, 403)


# ---------------------------------------------------------------------------
# PDF note parser — tested directly (public function)
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestPdfNoteParser:
    """
    parse_note_pdf(file_bytes) in catalog/note_pdf_parser.py.
    Uses pdfplumber internally — mocked here to avoid real PDF dependency.
    """

    def test_parser_returns_dict_or_list(self):
        """Smoke test: parser returns a structured result, not None."""
        from catalog.note_pdf_parser import parse_note_pdf

        with patch('catalog.note_pdf_parser.pdfplumber') as mock_plumber:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = 'Introduction to Waves\nWaves are...'
            mock_page.extract_words.return_value = [
                {'text': 'Introduction', 'top': 10, 'bottom': 20}
            ]
            mock_pdf = MagicMock()
            mock_pdf.pages = [mock_page]
            mock_pdf.__enter__ = lambda s: mock_pdf
            mock_pdf.__exit__ = MagicMock(return_value=False)
            mock_plumber.open.return_value = mock_pdf

            result = parse_note_pdf(b'fake-pdf-bytes')

        assert result is not None

    def test_parser_handles_empty_page(self):
        """Parser must not crash when a page has no extractable text."""
        from catalog.note_pdf_parser import parse_note_pdf

        with patch('catalog.note_pdf_parser.pdfplumber') as mock_plumber:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = None
            mock_page.extract_words.return_value = []
            mock_pdf = MagicMock()
            mock_pdf.pages = [mock_page]
            mock_pdf.__enter__ = lambda s: mock_pdf
            mock_pdf.__exit__ = MagicMock(return_value=False)
            mock_plumber.open.return_value = mock_pdf

            result = parse_note_pdf(b'fake-pdf-bytes')
            assert result is not None

    def test_parser_handles_multi_page_pdf(self):
        """Parser processes all pages without crashing."""
        from catalog.note_pdf_parser import parse_note_pdf

        with patch('catalog.note_pdf_parser.pdfplumber') as mock_plumber:
            pages = []
            for i in range(5):
                p = MagicMock()
                p.extract_text.return_value = f'Page {i} content about waves.'
                p.extract_words.return_value = [
                    {'text': f'Page{i}', 'top': 10, 'bottom': 20}
                ]
                pages.append(p)

            mock_pdf = MagicMock()
            mock_pdf.pages = pages
            mock_pdf.__enter__ = lambda s: mock_pdf
            mock_pdf.__exit__ = MagicMock(return_value=False)
            mock_plumber.open.return_value = mock_pdf

            result = parse_note_pdf(b'fake-pdf-bytes')
            assert result is not None


# ---------------------------------------------------------------------------
# Upload notes view (PDF upload endpoint)
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestUploadNotesView:

    def test_upload_notes_page_renders_for_admin(self, client, admin_user):
        client.force_login(admin_user)
        url = reverse('teacher:upload_notes')
        response = client.get(url)
        assert response.status_code == 200

    def test_upload_notes_requires_login(self, client):
        url = reverse('teacher:upload_notes')
        response = client.get(url)
        assert response.status_code == 302

    def test_non_admin_cannot_access_upload_notes(self, client, student):
        client.force_login(student)
        url = reverse('teacher:upload_notes')
        response = client.get(url)
        assert response.status_code in (302, 403)


# ---------------------------------------------------------------------------
# Past paper upload view
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestUploadPastPaperView:

    def test_upload_past_paper_renders_for_admin(self, client, admin_user):
        client.force_login(admin_user)
        url = reverse('teacher:upload_past_paper')
        response = client.get(url)
        assert response.status_code == 200

    def test_upload_past_paper_requires_login(self, client):
        url = reverse('teacher:upload_past_paper')
        response = client.get(url)
        assert response.status_code == 302

    def test_non_admin_cannot_upload_past_paper(self, client, student):
        client.force_login(student)
        url = reverse('teacher:upload_past_paper')
        response = client.get(url)
        assert response.status_code in (302, 403)


# ---------------------------------------------------------------------------
# Difficulty extraction — helpers and fixtures
# ---------------------------------------------------------------------------

from tests.conftest import (
    ExamBoardFactory, ExamSeriesFactory, SubjectFactory, TopicFactory,
)


def _make_docx_bytes(paragraphs):
    """Build a minimal in-memory DOCX from a list of paragraph strings."""
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _upload_docx(client, paragraphs):
    """
    POST a DOCX to the upload_docx view with overwrite enabled.
    Returns (response, error_message) where error_message is None on success.
    """
    docx_buf      = _make_docx_bytes(paragraphs)
    docx_buf.name = 'test.docx'
    response = client.post(
        reverse('teacher:upload_docx'),
        {'docx_file': docx_buf, 'overwrite': 'on'},
    )
    # Extract any error from template context for clearer test failure messages
    error = None
    if hasattr(response, 'context') and response.context and 'error' in response.context:
        error = response.context['error']
    return response, error


@pytest.fixture
def diff_subject(db):
    """Subject whose name matches the 'Subject:' header line in test DOCXs."""
    return SubjectFactory(name='Mathematics')


@pytest.fixture
def diff_board(db):
    """ExamBoard whose abbreviation matches the 'Exam:' header line in test DOCXs."""
    return ExamBoardFactory(name='WAEC', abbreviation='WAEC')


@pytest.fixture
def diff_series(db, diff_subject, diff_board):
    return ExamSeriesFactory(
        subject=diff_subject, exam_board=diff_board,
        year=2024, sitting='MAY_JUNE',
    )


@pytest.fixture
def diff_topic(db, diff_subject):
    """Topic name must match the 'Topic:' line in OBJ test DOCXs exactly."""
    return TopicFactory(name='Basic Arithmetic', subject=diff_subject)


@pytest.fixture
def diff_topic_theory(db, diff_subject):
    return TopicFactory(name='Forces and Motion', subject=diff_subject)


@pytest.fixture
def diff_topic_cell(db, diff_subject):
    return TopicFactory(name='Cell Biology', subject=diff_subject)


@pytest.fixture
def diff_topic_gas(db, diff_subject):
    return TopicFactory(name='Gas Laws', subject=diff_subject)


# ---------------------------------------------------------------------------
# _normalise_difficulty — pure unit tests, no DB needed
# ---------------------------------------------------------------------------

class TestNormaliseDifficulty:
    """
    _normalise_difficulty is a module-level helper in teacher/views.py.
    No Django DB needed — import and call directly.
    """

    @pytest.fixture(autouse=True)
    def _import_fn(self):
        from teacher.views import _normalise_difficulty
        self.fn = _normalise_difficulty

    def test_easy_variants(self):
        assert self.fn('Easy')    == 'EASY'
        assert self.fn('easy')    == 'EASY'
        assert self.fn('EASY')    == 'EASY'
        assert self.fn('  Easy ') == 'EASY'

    def test_medium_variants(self):
        assert self.fn('Medium') == 'MEDIUM'
        assert self.fn('medium') == 'MEDIUM'
        assert self.fn('MEDIUM') == 'MEDIUM'

    def test_hard_variants(self):
        assert self.fn('Hard') == 'HARD'
        assert self.fn('hard') == 'HARD'
        assert self.fn('HARD') == 'HARD'

    def test_unrecognised_returns_none(self):
        assert self.fn('Difficult') is None
        assert self.fn('1')         is None
        assert self.fn('')          is None
        assert self.fn('★★★')       is None
        assert self.fn('Tricky')    is None


# ---------------------------------------------------------------------------
# OBJ difficulty extraction — through the upload_docx view
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestDocxDifficultyExtractionObj:
    """
    Strategy: bypass _parse_docx entirely by mocking it to return a
    controlled parsed result. This avoids pandoc/DOCX format issues in CI
    and tests only the upload_docx view's handling of difficulty data.
    """

    def _mock_parse_result(self, subject, board, year, difficulty=None):
        """Return (header, questions) as _parse_docx would for a single OBJ question."""
        header = {
            'subject':    subject.name,
            'exam':       board.abbreviation,
            'year':       str(year),
            'sitting':    'MAY_JUNE',
            'paper_type': 'OBJ',
        }
        questions = [{
            'number':              1,
            'content':             '<p>What is 2 + 2?</p>',
            'content_after_image': '',
            'image_bytes':         None,
            'image_ext':           None,
            'choices': [
                {'label': 'A', 'text': '3', 'is_correct': False, 'explanation': ''},
                {'label': 'B', 'text': '4', 'is_correct': True,  'explanation': 'Add the numbers.'},
                {'label': 'C', 'text': '5', 'is_correct': False, 'explanation': ''},
                {'label': 'D', 'text': '6', 'is_correct': False, 'explanation': ''},
            ],
            'answer':      'B',
            'explanation': 'Add the numbers.',
            'difficulty':  difficulty,
            'topics':      [],
            'marks':       1,
        }]
        return header, questions

    def _upload(self, client, admin_user, subject, board, series, difficulty=None):
        client.force_login(admin_user)
        mock_result = self._mock_parse_result(subject, board, series.year, difficulty)
        with patch('teacher.views._parse_docx', return_value=mock_result):
            buf = io.BytesIO(b'fake-docx')
            buf.name = 'test.docx'
            return client.post(
                reverse('teacher:upload_docx'),
                {'docx_file': buf, 'overwrite': 'on'},
            )

    def test_difficulty_easy_stored_on_question(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'EASY')
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None, "Question was not created"
        assert q.difficulty == 'EASY'

    def test_difficulty_medium_stored(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'MEDIUM')
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None
        assert q.difficulty == 'MEDIUM'

    def test_difficulty_hard_stored(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'HARD')
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None
        assert q.difficulty == 'HARD'

    def test_missing_difficulty_leaves_field_null(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        """No difficulty in parsed data → Question.difficulty is None."""
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, None)
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None
        assert q.difficulty is None

    def test_difficulty_not_in_question_content(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        """Difficulty value must not appear in question.content."""
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'MEDIUM')
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None
        assert 'MEDIUM'     not in q.content
        assert 'Difficulty' not in q.content

    def test_difficulty_not_in_explanation(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        """Difficulty value must not appear in Choice.explanation."""
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'HARD')
        from catalog.models import Choice
        correct = Choice.objects.filter(
            question__exam_series=diff_series,
            question__question_number=1,
            is_correct=True,
        ).first()
        assert correct is not None
        assert 'HARD' not in (correct.explanation or '')

    def test_overwrite_updates_difficulty(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        """Re-uploading with Overwrite must update an existing question's difficulty."""
        from catalog.models import Question

        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'EASY')
        q = Question.objects.get(exam_series=diff_series, question_number=1)
        assert q.difficulty == 'EASY'

        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'HARD')
        q.refresh_from_db()
        assert q.difficulty == 'HARD'


# ---------------------------------------------------------------------------
# THEORY difficulty extraction — through the upload_docx view
# ---------------------------------------------------------------------------

@pytest.mark.django_db
class TestDocxDifficultyExtractionTheory:
    """
    Same mock strategy as TestDocxDifficultyExtractionObj — patches _parse_docx
    to return a controlled theory question dict with/without difficulty.
    """

    def _mock_parse_result(self, subject, board, year, difficulty=None):
        header = {
            'subject':    subject.name,
            'exam':       board.abbreviation,
            'year':       str(year),
            'sitting':    'MAY_JUNE',
            'paper_type': 'THEORY',
        }
        questions = [{
            'number':          1,
            'content':         '<p>Explain Newton\'s second law.</p>',
            'content_after_image': '',
            'image_bytes':     None,
            'image_ext':       None,
            'image_width_px':  None,
            'image_height_px': None,
            'choices':         [],
            'answer':          '',
            'theory_answer':   '<p>F = ma. Force equals mass times acceleration.</p>',
            'marking_guide':   '<p>1 mark for F = ma.</p>',
            'video_url':       None,
            'difficulty':      difficulty,
            'topics':          [],
            'marks':           2,
        }]
        return header, questions

    def _upload(self, client, admin_user, subject, board, series, difficulty=None):
        client.force_login(admin_user)
        mock_result = self._mock_parse_result(subject, board, series.year, difficulty)
        with patch('teacher.views._parse_docx', return_value=mock_result):
            buf = io.BytesIO(b'fake-docx')
            buf.name = 'test.docx'
            return client.post(
                reverse('teacher:upload_docx'),
                {'docx_file': buf, 'overwrite': 'on'},
            )

    def test_theory_difficulty_stored(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'MEDIUM')
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None, "Theory question was not created"
        assert q.difficulty == 'MEDIUM'

    def test_theory_difficulty_not_in_answer_content(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        """Difficulty value must not appear in TheoryAnswer.content."""
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, 'EASY')
        from catalog.models import TheoryAnswer
        ta = TheoryAnswer.objects.filter(
            question__exam_series=diff_series,
            question__question_number=1,
        ).first()
        assert ta is not None
        assert 'EASY'       not in ta.content
        assert 'Difficulty' not in ta.content

    def test_theory_missing_difficulty_leaves_null(
        self, client, admin_user, diff_subject, diff_board, diff_series
    ):
        """Theory question without difficulty → Question.difficulty is None."""
        self._upload(client, admin_user, diff_subject, diff_board, diff_series, None)
        from catalog.models import Question
        q = Question.objects.filter(
            exam_series=diff_series, question_number=1
        ).first()
        assert q is not None
        assert q.difficulty is None